from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO, Union

import pdfplumber
import pytesseract
from docx import Document
from PIL import Image, ImageOps

SourceLike = Union[str, Path, bytes, bytearray, BinaryIO, Image.Image]

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}
DOCX_EXTENSIONS = {".docx"}
TEXT_EXTENSIONS = {".txt", ".md", ".rst"}
PDF_EXTENSIONS = {".pdf"}


def _normalize_text(text: str) -> str:
    """Normalize whitespace while keeping blank lines intact."""
    lines = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        lines.append(re.sub(r"\s+", " ", line).strip())
    return "\n".join(lines).strip()


def _rewind_if_possible(source: object) -> None:
    if hasattr(source, "seek"):
        try:
            source.seek(0)
        except Exception:
            pass


def _extension_from_name(name: str | None) -> str:
    if not name:
        return ""
    return Path(name).suffix.lower()


def _coerce_image(source: SourceLike) -> Image.Image:
    if isinstance(source, Image.Image):
        return source

    if isinstance(source, (bytes, bytearray)):
        with Image.open(io.BytesIO(source)) as image:
            return image.copy()

    if isinstance(source, Path):
        with Image.open(source) as image:
            return image.copy()

    if isinstance(source, str):
        with Image.open(source) as image:
            return image.copy()

    if hasattr(source, "read"):
        _rewind_if_possible(source)
        with Image.open(source) as image:
            return image.copy()

    raise TypeError(f"Unsupported image source type: {type(source)!r}")


def _extract_text_from_docx_source(source: SourceLike) -> str:
    if isinstance(source, (bytes, bytearray)):
        document = Document(io.BytesIO(source))
    else:
        if hasattr(source, "seek"):
            _rewind_if_possible(source)
        document = Document(source)

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()
            )
            if row_text:
                chunks.append(row_text)

    return _normalize_text("\n".join(chunks))


def extract_text_from_image(source: SourceLike, language: str = "eng") -> str:
    """Run OCR on a single image and return normalized text."""
    image = _coerce_image(source)
    image = ImageOps.exif_transpose(image).convert("L")
    image = ImageOps.autocontrast(image)

    try:
        text = pytesseract.image_to_string(image, lang=language, config="--oem 3 --psm 6")
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError(
            "Tesseract OCR executable was not found. Install Tesseract or set "
            "pytesseract.pytesseract.tesseract_cmd before calling OCR functions."
        ) from exc

    return _normalize_text(text)


def extract_text_from_pdf(source: SourceLike, language: str = "eng") -> str:
    """Extract embedded text from a PDF, with OCR fallback for scanned pages."""
    pdf_source: object = source
    if isinstance(source, (bytes, bytearray)):
        pdf_source = io.BytesIO(source)
    elif isinstance(source, Path):
        pdf_source = str(source)
    elif isinstance(source, str):
        pdf_source = source
    elif hasattr(source, "seek"):
        _rewind_if_possible(source)

    chunks: list[str] = []
    with pdfplumber.open(pdf_source) as pdf:
        for page in pdf.pages:
            page_text = _normalize_text(page.extract_text() or "")
            if len(page_text.strip()) >= 20:
                chunks.append(page_text)
                continue

            if page_text.strip():
                chunks.append(page_text)
                continue

            try:
                page_image = page.to_image(resolution=300)
                original = getattr(page_image, "original", None)
                if original is not None:
                    ocr_text = extract_text_from_image(original, language=language)
                    if ocr_text.strip():
                        chunks.append(ocr_text)
            except Exception:
                continue

    return _normalize_text("\n\n".join(chunk for chunk in chunks if chunk.strip()))


def extract_text_from_docx(source: SourceLike) -> str:
    """Extract text from a DOCX file, including table content."""
    return _extract_text_from_docx_source(source)


def extract_text_from_bytes(
    data: bytes | bytearray,
    filename: str | None = None,
    language: str = "eng",
) -> str:
    """Extract text from raw bytes by inferring the file type when possible."""
    suffix = _extension_from_name(filename)
    payload = bytes(data)

    if suffix in TEXT_EXTENSIONS:
        return _normalize_text(payload.decode("utf-8", errors="ignore"))

    if suffix in IMAGE_EXTENSIONS:
        return extract_text_from_image(payload, language=language)

    if suffix in DOCX_EXTENSIONS:
        return extract_text_from_docx(payload)

    if suffix == PDF_EXTENSION:
        return extract_text_from_pdf(payload, language=language)

    if payload.startswith(b"%PDF"):
        return extract_text_from_pdf(io.BytesIO(payload), language=language)

    if payload.startswith(b"PK"):
        try:
            return extract_text_from_docx(payload)
        except Exception:
            pass

    try:
        return extract_text_from_image(payload, language=language)
    except Exception:
        pass

    try:
        return _normalize_text(payload.decode("utf-8"))
    except UnicodeDecodeError as exc:
        raise ValueError(
            "Unsupported binary source. Provide a PDF, DOCX, image, or UTF-8 text file."
        ) from exc


def extract_text(
    source: SourceLike,
    filename: str | None = None,
    language: str = "eng",
) -> str:
    """Extract text from PDF, image, DOCX, or raw text sources.

    Strings are treated as file paths when they exist on disk; otherwise they are
    treated as raw text.
    """
    if isinstance(source, Image.Image):
        return extract_text_from_image(source, language=language)

    if isinstance(source, (bytes, bytearray)):
        return extract_text_from_bytes(source, filename=filename, language=language)

    if isinstance(source, Path):
        if not source.exists():
            raise FileNotFoundError(f"File does not exist: {source}")
        suffix = source.suffix.lower()
        if suffix in PDF_EXTENSIONS:
            return extract_text_from_pdf(source, language=language)
        if suffix in IMAGE_EXTENSIONS:
            return extract_text_from_image(source, language=language)
        if suffix in DOCX_EXTENSIONS:
            return extract_text_from_docx(source)
        if suffix in TEXT_EXTENSIONS:
            return _normalize_text(source.read_text(encoding="utf-8", errors="ignore"))
        raise ValueError(f"Unsupported file type: {suffix or '<none>'}")

    if isinstance(source, str):
        path = Path(source)
        if path.exists():
            return extract_text(path, language=language)
        return _normalize_text(source)

    if hasattr(source, "read"):
        _rewind_if_possible(source)
        inferred_name = filename or getattr(source, "name", None)
        suffix = _extension_from_name(inferred_name)

        if suffix in PDF_EXTENSIONS:
            return extract_text_from_pdf(source, language=language)
        if suffix in IMAGE_EXTENSIONS:
            return extract_text_from_image(source, language=language)
        if suffix in DOCX_EXTENSIONS:
            return extract_text_from_docx(source)
        if suffix in TEXT_EXTENSIONS:
            data = source.read()
            _rewind_if_possible(source)
            if isinstance(data, str):
                return _normalize_text(data)
            return _normalize_text(data.decode("utf-8", errors="ignore"))

        data = source.read()
        _rewind_if_possible(source)
        if isinstance(data, str):
            return _normalize_text(data)
        return extract_text_from_bytes(data, filename=inferred_name, language=language)

    raise TypeError(f"Unsupported source type: {type(source)!r}")


# Backwards-compatible alias for callers that prefer an OCR-style name.
ocr_document = extract_text


__all__ = [
    "extract_text",
    "extract_text_from_bytes",
    "extract_text_from_docx",
    "extract_text_from_image",
    "extract_text_from_pdf",
    "ocr_document",
]
